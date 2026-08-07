from __future__ import annotations

import argparse
import json
import os
import shutil
import subprocess
import sys
import tempfile
import zipfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from lxml import etree


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}
W = f"{{{W_NS}}}"
BLACK = "000000"
WHITE = "FFFFFF"
HEADING_LATIN = "SimHei"
HEADING_EAST_ASIA = "黑体"
BODY_LATIN = "Microsoft YaHei"
BODY_EAST_ASIA = "微软雅黑"
TEXT_TAGS = {f"{W}t", f"{W}instrText", f"{W}delText"}
HEADING_STYLE_KEYS = ("title", "subtitle", "heading", "标题")


@dataclass(frozen=True)
class StyleProfile:
    heading_latin: str = HEADING_LATIN
    heading_east_asia: str = HEADING_EAST_ASIA
    body_latin: str = BODY_LATIN
    body_east_asia: str = BODY_EAST_ASIA
    heading_size_pt: float | None = None
    body_size_pt: float | None = None
    page_number_size_pt: float = 10.0

    @property
    def mode(self) -> str:
        return "default" if self == StyleProfile() else "custom"

    def to_dict(self) -> dict:
        return {
            "mode": self.mode,
            "heading_font": {
                "latin": self.heading_latin,
                "east_asia": self.heading_east_asia,
            },
            "body_font": {
                "latin": self.body_latin,
                "east_asia": self.body_east_asia,
            },
            "heading_size_pt": self.heading_size_pt or "preserve",
            "body_size_pt": self.body_size_pt or "preserve",
            "page_number_size_pt": self.page_number_size_pt,
        }


def is_heading_style(name: str | None, style_id: str | None = None) -> bool:
    value = f"{name or ''} {style_id or ''}".lower()
    return any(key in value for key in HEADING_STYLE_KEYS)


def get_or_add(parent, tag: str, *, first: bool = False):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        if first:
            parent.insert(0, child)
        else:
            parent.append(child)
    return child


def font_names(profile: StyleProfile, *, heading: bool) -> tuple[str, str]:
    if heading:
        return profile.heading_latin, profile.heading_east_asia
    return profile.body_latin, profile.body_east_asia


def profile_size(profile: StyleProfile, *, heading: bool) -> float | None:
    return profile.heading_size_pt if heading else profile.body_size_pt


def set_rpr_size(rpr, size_pt: float) -> None:
    half_points = str(round(size_pt * 2))
    size = get_or_add(rpr, "w:sz")
    size.set(qn("w:val"), half_points)
    size_cs = get_or_add(rpr, "w:szCs")
    size_cs.set(qn("w:val"), half_points)


def set_rpr_font_and_color(
    rpr,
    *,
    heading: bool,
    profile: StyleProfile,
    size_pt: float | None = None,
) -> None:
    latin, east_asia = font_names(profile, heading=heading)
    rfonts = get_or_add(rpr, "w:rFonts", first=True)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:cs"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    for attr in ("asciiTheme", "hAnsiTheme", "csTheme", "eastAsiaTheme", "hint"):
        rfonts.attrib.pop(qn(f"w:{attr}"), None)

    color = get_or_add(rpr, "w:color")
    color.set(qn("w:val"), BLACK)
    for attr in ("themeColor", "themeTint", "themeShade"):
        color.attrib.pop(qn(f"w:{attr}"), None)

    selected_size = profile_size(profile, heading=heading) if size_pt is None else size_pt
    if selected_size is not None:
        set_rpr_size(rpr, selected_size)


def normalize_styles(doc: Document, profile: StyleProfile) -> None:
    for style in doc.styles:
        if style.type not in (WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER):
            continue
        heading = is_heading_style(style.name, style.style_id)
        latin, _ = font_names(profile, heading=heading)
        style.font.name = latin
        style.font.color.rgb = RGBColor(0, 0, 0)
        size_pt = profile_size(profile, heading=heading)
        if size_pt is not None:
            style.font.size = Pt(size_pt)
        rpr = style._element.get_or_add_rPr()
        set_rpr_font_and_color(rpr, heading=heading, profile=profile)


def set_cell_margins(tc_pr) -> None:
    tc_mar = get_or_add(tc_pr, "w:tcMar")
    for name, value in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
        margin = get_or_add(tc_mar, f"w:{name}")
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")


def set_paragraph_centered(p) -> None:
    p_pr = get_or_add(p, "w:pPr", first=True)
    jc = get_or_add(p_pr, "w:jc")
    jc.set(qn("w:val"), "center")


def set_run_bold(r) -> None:
    r_pr = get_or_add(r, "w:rPr", first=True)
    for tag in ("w:b", "w:bCs"):
        bold = get_or_add(r_pr, tag)
        bold.set(qn("w:val"), "1")


def normalize_table_xml(tbl) -> None:
    tbl_pr = get_or_add(tbl, "w:tblPr", first=True)
    borders = get_or_add(tbl_pr, "w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = get_or_add(borders, f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), BLACK)
        for attr in ("themeColor", "themeTint", "themeShade"):
            border.attrib.pop(qn(f"w:{attr}"), None)

    rows = [child for child in tbl if child.tag == qn("w:tr")]
    for row_index, row in enumerate(rows):
        tr_pr = get_or_add(row, "w:trPr", first=True)
        height = tr_pr.find(qn("w:trHeight"))
        if height is not None and height.get(qn("w:hRule")) == "exact":
            height.set(qn("w:hRule"), "atLeast")
        if row_index == 0:
            header = get_or_add(tr_pr, "w:tblHeader")
            header.set(qn("w:val"), "true")

        for cell in (child for child in row if child.tag == qn("w:tc")):
            tc_pr = get_or_add(cell, "w:tcPr", first=True)
            old_cell_borders = tc_pr.find(qn("w:tcBorders"))
            if old_cell_borders is not None:
                tc_pr.remove(old_cell_borders)
            shading = get_or_add(tc_pr, "w:shd")
            shading.set(qn("w:val"), "clear")
            shading.set(qn("w:color"), "auto")
            shading.set(qn("w:fill"), WHITE)
            v_align = get_or_add(tc_pr, "w:vAlign")
            v_align.set(qn("w:val"), "center")
            set_cell_margins(tc_pr)

            if row_index == 0:
                for paragraph in cell.iter(qn("w:p")):
                    set_paragraph_centered(paragraph)
                    for run in paragraph.iter(qn("w:r")):
                        set_run_bold(run)


def normalize_tables(doc: Document) -> None:
    for table in doc.tables:
        normalize_table_xml(table._tbl)


def story_key(story) -> str:
    return str(story.part.partname)


def clear_story(story):
    root = story._element
    for child in list(root):
        root.remove(child)
    root.append(OxmlElement("w:p"))
    return story.paragraphs[0]


def footer_has_page_field(story) -> bool:
    instructions = "".join(
        node.text or "" for node in story._element.iter(qn("w:instrText"))
    )
    return "PAGE" in instructions.upper()


def add_page_number(paragraph, profile: StyleProfile) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    def append_run_with(child):
        run = paragraph.add_run()
        rpr = run._element.get_or_add_rPr()
        set_rpr_font_and_color(
            rpr,
            heading=False,
            profile=profile,
            size_pt=profile.page_number_size_pt,
        )
        run._element.append(child)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    append_run_with(begin)

    instruction = OxmlElement("w:instrText")
    instruction.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instruction.text = " PAGE "
    append_run_with(instruction)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    append_run_with(separate)

    cached = OxmlElement("w:t")
    cached.text = "1"
    append_run_with(cached)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    append_run_with(end)


def normalize_headers_and_footers(
    doc: Document, page_numbers: str, profile: StyleProfile
) -> None:
    source_has_page = False
    for section in doc.sections:
        source_has_page = source_has_page or footer_has_page_field(section.footer)
        if section.different_first_page_header_footer:
            source_has_page = source_has_page or footer_has_page_field(
                section.first_page_footer
            )
        if getattr(doc.settings, "odd_and_even_pages_header_footer", False):
            source_has_page = source_has_page or footer_has_page_field(
                section.even_page_footer
            )

    use_page_numbers = (
        page_numbers == "always"
        or (page_numbers == "preserve" and source_has_page)
    )

    headers = {}
    all_footers = {}
    numbered_footers = {}
    first_page_footers = set()
    odd_even = getattr(doc.settings, "odd_and_even_pages_header_footer", False)

    for section in doc.sections:
        for story in (section.header, section.first_page_header, section.even_page_header):
            headers[story_key(story)] = story

        main_footer = section.footer
        all_footers[story_key(main_footer)] = main_footer
        numbered_footers[story_key(main_footer)] = main_footer

        even_footer = section.even_page_footer
        all_footers[story_key(even_footer)] = even_footer
        if odd_even:
            numbered_footers[story_key(even_footer)] = even_footer

        first_footer = section.first_page_footer
        all_footers[story_key(first_footer)] = first_footer
        if section.different_first_page_header_footer:
            first_page_footers.add(story_key(first_footer))
        else:
            numbered_footers[story_key(first_footer)] = first_footer

    for story in headers.values():
        clear_story(story)
    footer_paragraphs = {
        key: clear_story(story) for key, story in all_footers.items()
    }

    if use_page_numbers:
        for key in numbered_footers:
            if key not in first_page_footers:
                add_page_number(footer_paragraphs[key], profile)


def paragraph_is_heading_xml(paragraph) -> bool:
    style = paragraph.find("w:pPr/w:pStyle", namespaces=NS)
    style_id = style.get(f"{W}val") if style is not None else ""
    return is_heading_style(style_id)


def normalize_all_wordprocessing_runs(path: Path, profile: StyleProfile) -> None:
    temp_path = path.with_name(f".{path.name}.formatting.tmp")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(
        temp_path, "w", zipfile.ZIP_DEFLATED
    ) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                try:
                    root = etree.fromstring(data)
                except etree.XMLSyntaxError:
                    target.writestr(item, data)
                    continue

                for table in root.iter(f"{W}tbl"):
                    normalize_table_xml(table)

                for paragraph in root.iter(f"{W}p"):
                    heading = paragraph_is_heading_xml(paragraph)
                    for run in paragraph.iter(f"{W}r"):
                        if not any(child.tag in TEXT_TAGS for child in run):
                            continue
                        rpr = run.find(f"{W}rPr")
                        if rpr is None:
                            rpr = etree.Element(f"{W}rPr")
                            run.insert(0, rpr)
                        set_rpr_font_and_color(
                            rpr,
                            heading=heading,
                            profile=profile,
                            size_pt=(
                                profile.page_number_size_pt
                                if item.filename.startswith("word/footer")
                                else None
                            ),
                        )

                if item.filename == "word/numbering.xml":
                    for rpr in root.iter(f"{W}rPr"):
                        set_rpr_font_and_color(
                            rpr, heading=False, profile=profile
                        )

                data = etree.tostring(
                    root, xml_declaration=True, encoding="UTF-8", standalone=True
                )
            target.writestr(item, data)
    os.replace(temp_path, path)


def find_soffice() -> str | None:
    found = shutil.which("soffice") or shutil.which("libreoffice")
    if found:
        return found
    candidates = [
        Path(r"C:\Program Files\LibreOffice\program\soffice.exe"),
        Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"),
        Path("/usr/bin/libreoffice"),
        Path("/usr/bin/soffice"),
    ]
    return str(next((path for path in candidates if path.exists()), "")) or None


def convert_with_soffice(source: Path, output: Path) -> bool:
    executable = find_soffice()
    if not executable:
        return False
    output.parent.mkdir(parents=True, exist_ok=True)
    if output.suffix.lower() == ".docx":
        conversion = "docx"
    elif output.suffix.lower() == ".doc":
        conversion = 'doc:"MS Word 97"'
    else:
        raise ValueError(f"Unsupported conversion target: {output.suffix}")
    with tempfile.TemporaryDirectory(prefix="word_writer_lo_") as profile:
        profile_uri = Path(profile).resolve().as_uri()
        command = [
            executable,
            "--headless",
            f"-env:UserInstallation={profile_uri}",
            "--convert-to",
            conversion,
            "--outdir",
            str(output.parent),
            str(source),
        ]
        completed = subprocess.run(command, capture_output=True, text=True)
        generated = output.parent / f"{source.stem}{output.suffix}"
        if completed.returncode != 0 or not generated.exists():
            return False
        if generated.resolve() != output.resolve():
            os.replace(generated, output)
    return True


def convert_with_word_com(source: Path, output: Path) -> bool:
    if sys.platform != "win32":
        return False
    try:
        import win32com.client  # type: ignore
    except ImportError:
        return False

    file_format = 16 if output.suffix.lower() == ".docx" else 0
    word = None
    document = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(
            str(source.resolve()), ReadOnly=True, AddToRecentFiles=False
        )
        document.SaveAs2(str(output.resolve()), FileFormat=file_format)
        return output.exists()
    except Exception:
        return False
    finally:
        if document is not None:
            document.Close(False)
        if word is not None:
            word.Quit()


def convert_with_word_powershell(source: Path, output: Path) -> bool:
    if sys.platform != "win32":
        return False

    script = r'''
param(
    [Parameter(Mandatory = $true)][string]$SourcePath,
    [Parameter(Mandatory = $true)][string]$OutputPath,
    [Parameter(Mandatory = $true)][int]$FileFormat
)

$word = $null
$document = $null
try {
    $word = New-Object -ComObject Word.Application
    $word.Visible = $false
    $word.DisplayAlerts = 0
    $document = $word.Documents.Open($SourcePath, $false, $true, $false)
    $document.SaveAs2($OutputPath, $FileFormat)
    if (-not (Test-Path -LiteralPath $OutputPath)) { exit 1 }
}
catch {
    [Console]::Error.WriteLine($_.Exception.Message)
    exit 1
}
finally {
    if ($null -ne $document) { $document.Close($false) }
    if ($null -ne $word) { $word.Quit() }
}
'''
    file_format = 16 if output.suffix.lower() == ".docx" else 0
    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".ps1", encoding="utf-8", delete=False
    ) as script_file:
        script_file.write(script)
        script_path = Path(script_file.name)
    try:
        completed = subprocess.run(
            [
                "powershell.exe",
                "-NoProfile",
                "-NonInteractive",
                "-File",
                str(script_path),
                "-SourcePath",
                str(source.resolve()),
                "-OutputPath",
                str(output.resolve()),
                "-FileFormat",
                str(file_format),
            ],
            capture_output=True,
            text=True,
        )
        return completed.returncode == 0 and output.exists()
    finally:
        script_path.unlink(missing_ok=True)


def convert_word_file(source: Path, output: Path) -> None:
    if convert_with_soffice(source, output):
        return
    if convert_with_word_com(source, output):
        return
    if convert_with_word_powershell(source, output):
        return
    raise RuntimeError(
        "Legacy .doc conversion requires LibreOffice/soffice or Microsoft Word. "
        "Convert the file to .docx manually, then retry."
    )


def style_font_matches(
    rfonts, *, heading: bool, profile: StyleProfile
) -> bool:
    if rfonts is None:
        return False
    latin, east_asia = font_names(profile, heading=heading)
    return (
        rfonts.get(f"{W}ascii") == latin
        and rfonts.get(f"{W}hAnsi") == latin
        and rfonts.get(f"{W}cs") == latin
        and rfonts.get(f"{W}eastAsia") == east_asia
    )


def style_size_matches(rpr, size_pt: float | None) -> bool:
    if size_pt is None:
        return True
    if rpr is None:
        return False
    expected = str(round(size_pt * 2))
    size = rpr.find(f"{W}sz")
    size_cs = rpr.find(f"{W}szCs")
    return (
        size is not None
        and size_cs is not None
        and size.get(f"{W}val") == expected
        and size_cs.get(f"{W}val") == expected
    )


def audit_docx(path: Path, profile: StyleProfile) -> dict:
    issues: list[str] = []
    counts = {
        "text_runs": 0,
        "tables": 0,
        "headers": 0,
        "footers": 0,
        "images": 0,
    }

    with zipfile.ZipFile(path, "r") as archive:
        names = set(archive.namelist())
        counts["images"] = sum(
            name.startswith("word/media/") and not name.endswith("/") for name in names
        )
        parsed = {}
        for name in names:
            if name.startswith("word/") and name.endswith(".xml"):
                try:
                    parsed[name] = etree.fromstring(archive.read(name))
                except etree.XMLSyntaxError as exc:
                    issues.append(f"{name}: invalid XML ({exc})")

        for name, root in parsed.items():
            for paragraph in root.iter(f"{W}p"):
                heading = paragraph_is_heading_xml(paragraph)
                for run in paragraph.iter(f"{W}r"):
                    if not any(child.tag in TEXT_TAGS for child in run):
                        continue
                    counts["text_runs"] += 1
                    rpr = run.find(f"{W}rPr")
                    color = rpr.find(f"{W}color") if rpr is not None else None
                    rfonts = rpr.find(f"{W}rFonts") if rpr is not None else None
                    if color is None or color.get(f"{W}val") != BLACK:
                        issues.append(f"{name}: text run without explicit black color")
                    if not style_font_matches(
                        rfonts, heading=heading, profile=profile
                    ):
                        role = "heading" if heading else "body"
                        issues.append(f"{name}: {role} run has a nonstandard font")
                    expected_size = (
                        profile.page_number_size_pt
                        if name.startswith("word/footer")
                        else profile_size(profile, heading=heading)
                    )
                    if not style_size_matches(rpr, expected_size):
                        role = "page-number" if name.startswith("word/footer") else (
                            "heading" if heading else "body"
                        )
                        issues.append(f"{name}: {role} run has a nonstandard size")

        document_root = parsed.get("word/document.xml")
        if document_root is not None:
            for table_index, table in enumerate(
                document_root.iter(f"{W}tbl"), start=1
            ):
                counts["tables"] += 1
                for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                    border = table.find(
                        f"w:tblPr/w:tblBorders/w:{edge}", namespaces=NS
                    )
                    if (
                        border is None
                        or border.get(f"{W}val") != "single"
                        or border.get(f"{W}color") != BLACK
                    ):
                        issues.append(
                            f"table {table_index}: {edge} border is not single black"
                        )

                rows = table.findall("w:tr", namespaces=NS)
                if not rows:
                    continue
                header = rows[0]
                header_flag = header.find("w:trPr/w:tblHeader", namespaces=NS)
                if header_flag is None:
                    issues.append(
                        f"table {table_index}: first row is not marked as header"
                    )

                for cell in table.iter(f"{W}tc"):
                    shading = cell.find("w:tcPr/w:shd", namespaces=NS)
                    if shading is None or shading.get(f"{W}fill") != WHITE:
                        issues.append(f"table {table_index}: cell fill is not white")
                    v_align = cell.find("w:tcPr/w:vAlign", namespaces=NS)
                    if v_align is None or v_align.get(f"{W}val") != "center":
                        issues.append(
                            f"table {table_index}: cell is not vertically centered"
                        )

                for paragraph in header.iter(f"{W}p"):
                    jc = paragraph.find("w:pPr/w:jc", namespaces=NS)
                    if jc is None or jc.get(f"{W}val") != "center":
                        issues.append(
                            f"table {table_index}: header text is not centered"
                        )
                    for run in paragraph.iter(f"{W}r"):
                        if not any(child.tag in TEXT_TAGS for child in run):
                            continue
                        bold = run.find("w:rPr/w:b", namespaces=NS)
                        if bold is None or bold.get(f"{W}val", "1") in (
                            "0",
                            "false",
                            "off",
                        ):
                            issues.append(
                                f"table {table_index}: header text is not bold"
                            )

        for name, root in parsed.items():
            if name.startswith("word/header"):
                counts["headers"] += 1
                visible = "".join(node.text or "" for node in root.iter(f"{W}t"))
                if visible.strip():
                    issues.append(f"{name}: header contains visible text")
            elif name.startswith("word/footer"):
                counts["footers"] += 1
                instructions = "".join(
                    node.text or "" for node in root.iter(f"{W}instrText")
                ).upper()
                visible = "".join(node.text or "" for node in root.iter(f"{W}t"))
                if "PAGE" in instructions:
                    if visible.strip() and not visible.strip().isdigit():
                        issues.append(f"{name}: footer contains text beyond page number")
                    page_paragraphs = [
                        p
                        for p in root.iter(f"{W}p")
                        if "PAGE"
                        in "".join(
                            node.text or ""
                            for node in p.iter(f"{W}instrText")
                        ).upper()
                    ]
                    for paragraph in page_paragraphs:
                        jc = paragraph.find("w:pPr/w:jc", namespaces=NS)
                        if jc is None or jc.get(f"{W}val") != "center":
                            issues.append(f"{name}: page number is not centered")
                elif visible.strip():
                    issues.append(f"{name}: footer contains non-page-number text")

    unique_issues = list(dict.fromkeys(issues))
    return {
        "file": str(path.resolve()),
        "status": "pass" if not unique_issues else "fail",
        "style_profile": profile.to_dict(),
        "counts": counts,
        "issues": unique_issues,
    }


def normalize_docx(
    source: Path,
    output: Path,
    page_numbers: str,
    profile: StyleProfile,
) -> dict:
    doc = Document(source)
    normalize_styles(doc, profile)
    normalize_tables(doc)
    normalize_headers_and_footers(doc, page_numbers, profile)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    normalize_all_wordprocessing_runs(output, profile)
    return audit_docx(output, profile)


def run_normalization(
    source: Path,
    output: Path,
    page_numbers: str,
    profile: StyleProfile,
) -> dict:
    source_suffix = source.suffix.lower()
    output_suffix = output.suffix.lower()
    if source_suffix not in (".doc", ".docx"):
        raise ValueError("Input must be a .doc or .docx file")
    if output_suffix not in (".doc", ".docx"):
        raise ValueError("Output must be a .doc or .docx file")
    if not source.exists():
        raise FileNotFoundError(source)

    with tempfile.TemporaryDirectory(prefix="word_writer_") as temp_dir:
        temp_root = Path(temp_dir)
        working_source = source
        if source_suffix == ".doc":
            working_source = temp_root / f"{source.stem}.docx"
            convert_word_file(source, working_source)

        if output_suffix == ".docx":
            report = normalize_docx(working_source, output, page_numbers, profile)
        else:
            normalized_docx = temp_root / f"{output.stem}.normalized.docx"
            report = normalize_docx(
                working_source, normalized_docx, page_numbers, profile
            )
            convert_word_file(normalized_docx, output)
            roundtrip_docx = temp_root / f"{output.stem}.roundtrip.docx"
            convert_word_file(output, roundtrip_docx)
            report = audit_docx(roundtrip_docx, profile)
            report["file"] = str(output.resolve())
            report["note"] = "Audit performed after DOC -> DOCX round-trip."
    return report


def point_size(value: str) -> float:
    size = float(value)
    if size < 1 or size > 200:
        raise argparse.ArgumentTypeError("size must be between 1 and 200 points")
    return size


def font_name(value: str) -> str:
    name = value.strip()
    if not name:
        raise argparse.ArgumentTypeError("font name cannot be empty")
    return name


def build_style_profile(args: argparse.Namespace) -> StyleProfile:
    default = StyleProfile()
    heading_latin = args.heading_font or default.heading_latin
    heading_east_asia = args.heading_font or default.heading_east_asia
    body_latin = args.body_font or default.body_latin
    body_east_asia = args.body_font or default.body_east_asia
    return StyleProfile(
        heading_latin=heading_latin,
        heading_east_asia=heading_east_asia,
        body_latin=body_latin,
        body_east_asia=body_east_asia,
        heading_size_pt=args.heading_size,
        body_size_pt=args.body_size,
        page_number_size_pt=args.page_number_size,
    )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Normalize .doc/.docx typography, tables, headers, and footers."
    )
    parser.add_argument("input", type=Path)
    parser.add_argument("--output", type=Path)
    parser.add_argument(
        "--page-numbers",
        choices=("preserve", "always", "none"),
        default="preserve",
    )
    parser.add_argument(
        "--heading-font",
        type=font_name,
        help="set one font name for all heading text slots",
    )
    parser.add_argument(
        "--body-font",
        type=font_name,
        help="set one font name for all body text slots",
    )
    parser.add_argument(
        "--heading-size",
        type=point_size,
        help="set one heading size in points; omit to preserve existing sizes",
    )
    parser.add_argument(
        "--body-size",
        type=point_size,
        help="set one body size in points; omit to preserve existing sizes",
    )
    parser.add_argument(
        "--page-number-size",
        type=point_size,
        default=10.0,
        help="set page-number size in points (default: 10)",
    )
    parser.add_argument(
        "--audit-only",
        action="store_true",
        help="Check an existing DOCX without modifying it.",
    )
    parser.add_argument(
        "--force",
        action="store_true",
        help="overwrite an existing output file; input and output must still differ",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    profile = build_style_profile(args)
    if args.audit_only:
        if args.input.suffix.lower() != ".docx":
            raise ValueError("--audit-only currently requires a .docx file")
        report = audit_docx(args.input, profile)
    else:
        if args.output is None:
            raise ValueError("--output is required unless --audit-only is used")
        if args.input.resolve() == args.output.resolve():
            raise ValueError("input and output paths must differ; normalize a copy")
        if args.output.exists() and not args.force:
            raise FileExistsError(
                f"output already exists: {args.output}; use a new path or pass --force"
            )
        report = run_normalization(
            args.input, args.output, args.page_numbers, profile
        )

    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0 if report["status"] == "pass" else 1


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(
            json.dumps(
                {"status": "error", "error": str(exc)},
                ensure_ascii=False,
                indent=2,
            ),
            file=sys.stderr,
        )
        raise SystemExit(2)
