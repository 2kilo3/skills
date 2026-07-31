from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


REQUIRED_FIELDS = (
    "title",
    "authors",
    "document_type",
    "source",
    "structure",
    "research_question",
    "methods",
    "results",
    "innovation",
    "reflection",
)

EXPECTED_TEMPLATE_SHA256 = "4824159895BDA6297DEF7DEFDCE79CD406D88BE40C6D9ED4AA9ABBEF680A387C"
NUMBER_PREFIX_RE = re.compile(r"^\s*\d+\s*[.．、]\s*")
LABEL_PREFIX_RE = re.compile(r"^[^，。；！？,.!?;：:\n]{1,20}[：:]")
STRUCTURE_TOP_RE = re.compile(
    r"^(?P<number>\d+)\.\s+(?P<title>[^：\n]+)：(?P<summary>\S.*)$"
)
STRUCTURE_SUB_RE = re.compile(
    r"^(?P<parent>\d+)\.(?P<number>\d+)\s+(?P<title>[^：\n]+)：(?P<summary>\S.*)$"
)
STRUCTURE_SUB_LEFT_INDENT = Cm(0.74)
STRUCTURE_SUB_HANGING_INDENT = Cm(-0.37)


def file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def resolve_template(skill_dir: Path) -> Path:
    template = skill_dir / "assets" / "reading-note-template.docx"
    if not template.is_file():
        raise FileNotFoundError(f"bundled template is missing: {template}")
    actual_hash = file_sha256(template)
    if actual_hash != EXPECTED_TEMPLATE_SHA256:
        raise ValueError(
            "bundled template SHA-256 does not match the locked template: "
            f"{template} ({actual_hash})"
        )
    return template


def normalize_lines(value: Any, *, separator: str | None = None) -> list[str]:
    if value is None:
        return []
    if isinstance(value, list):
        items = [str(item).strip() for item in value]
    else:
        text = str(value).replace("\r\n", "\n").replace("\r", "\n")
        items = [line.strip() for line in text.split("\n")]
    items = [item for item in items if item]
    if separator is not None and len(items) > 1:
        return [separator.join(items)]
    return items


def normalize_numbered_points(value: Any, *, field_name: str) -> list[str]:
    points = []
    for line in normalize_lines(value):
        point = NUMBER_PREFIX_RE.sub("", line).strip()
        if LABEL_PREFIX_RE.match(point):
            raise ValueError(
                f"{field_name} points must start directly with sentence content; "
                "do not use a short label followed by a colon"
            )
        points.append(point)
    if len(points) > 1:
        return [f"{index}. {point}" for index, point in enumerate(points, start=1)]
    return points


def normalize_structure(value: Any) -> list[str]:
    lines = normalize_lines(value)
    normalized: list[str] = []
    expected_top = 1
    current_top: int | None = None
    expected_sub = 1

    for line in lines:
        sub_match = STRUCTURE_SUB_RE.fullmatch(line)
        if sub_match:
            parent = int(sub_match.group("parent"))
            number = int(sub_match.group("number"))
            if current_top is None or parent != current_top:
                raise ValueError(
                    "structure subsection must follow its numbered top-level section"
                )
            if number != expected_sub:
                raise ValueError(
                    f"structure subsection under {parent}. must continue with "
                    f"{parent}.{expected_sub}"
                )
            expected_sub += 1
            normalized.append(line)
            continue

        top_match = STRUCTURE_TOP_RE.fullmatch(line)
        if not top_match:
            raise ValueError(
                "structure items must use Arabic numbering and a Chinese colon: "
                "'1. Real section title：one-sentence summary' or "
                "'3.1 Real subsection title：one-sentence summary'"
            )
        number = int(top_match.group("number"))
        if number != expected_top:
            raise ValueError(
                f"structure top-level numbering must be continuous and start with "
                f"{expected_top}."
            )
        current_top = number
        expected_top += 1
        expected_sub = 1
        normalized.append(line)

    return normalized


def set_run_font(run, *, size_pt: float, bold: bool = False) -> None:
    run.font.name = "微软雅黑"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(key), "微软雅黑")


def clear_cell(cell) -> None:
    tc = cell._tc
    for child in list(tc):
        if child.tag != qn("w:tcPr"):
            tc.remove(child)
    tc.append(OxmlElement("w:p"))


def fill_cell(
    cell,
    value: Any,
    *,
    bold_headings: bool = False,
    structure_hierarchy: bool = False,
) -> None:
    lines = normalize_lines(value)
    clear_cell(cell)
    if not lines:
        return

    first = True
    for line in lines:
        paragraph = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.2
        display_line = line
        if structure_hierarchy:
            if STRUCTURE_SUB_RE.fullmatch(line):
                # Match the teacher reference's "-" subpoint marker while using
                # real paragraph indentation so wrapped lines stay aligned.
                paragraph.paragraph_format.left_indent = STRUCTURE_SUB_LEFT_INDENT
                paragraph.paragraph_format.first_line_indent = STRUCTURE_SUB_HANGING_INDENT
                display_line = f"- {line}"
            else:
                paragraph.paragraph_format.left_indent = Cm(0)
                paragraph.paragraph_format.first_line_indent = Cm(0)
        run = paragraph.add_run(display_line)
        is_heading = bold_headings and line.rstrip() in {"创新：", "不足：", "创新:", "不足:"}
        set_run_font(run, size_pt=10, bold=is_heading)


def fill_label(cell, value: str) -> None:
    clear_cell(cell)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(value)
    set_run_font(run, size_pt=10.5, bold=True)


def load_content(path: Path) -> dict[str, Any]:
    data = json.loads(path.read_text(encoding="utf-8-sig"))
    if not isinstance(data, dict):
        raise ValueError("content JSON must contain one object")
    missing = [field for field in REQUIRED_FIELDS if not normalize_lines(data.get(field))]
    if missing:
        raise ValueError("missing required fields: " + ", ".join(missing))
    return data


def validate_template(doc: Document) -> None:
    if len(doc.sections) != 1 or len(doc.tables) != 1:
        raise ValueError("template must contain exactly one section and one table")
    table = doc.tables[0]
    if len(table.rows) != 13 or len(table.columns) != 5:
        raise ValueError("template table must be 13 rows by 5 columns")
    expected = {
        (1, 1): "题目",
        (2, 1): "作者",
        (2, 3): "文献类型",
        (3, 1): "影响因子",
        (3, 3): "文献来源",
        (4, 1): "原文链接",
        (5, 1): "关键词",
        (6, 1): "论文结构",
        (7, 1): "研究问题",
        (8, 1): "研究方法",
        (9, 1): "研究结果",
        (11, 1): "感悟心得",
        (12, 0): "备注",
    }
    for (row, column), label in expected.items():
        if table.cell(row, column).text.strip() != label:
            raise ValueError(f"unexpected template label at row {row + 1}, column {column + 1}")


def compact_trailing_body_paragraphs(doc: Document) -> None:
    """Keep one 1 pt empty paragraph after the table to prevent a blank last page."""
    body = doc._element.body
    table_seen = False
    trailing_paragraphs = []

    for child in list(body):
        if child.tag == qn("w:tbl"):
            table_seen = True
            continue
        if not table_seen or child.tag != qn("w:p"):
            continue
        if "".join(child.itertext()).strip():
            raise ValueError("template must not contain body text outside the table")
        trailing_paragraphs.append(child)

    if not trailing_paragraphs:
        paragraph_element = OxmlElement("w:p")
        body.insert(body.index(body.sectPr), paragraph_element)
        trailing_paragraphs.append(paragraph_element)

    for paragraph_element in trailing_paragraphs[1:]:
        body.remove(paragraph_element)

    paragraph = Paragraph(trailing_paragraphs[0], doc._body)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(1)
    paragraph.paragraph_format.page_break_before = False
    paragraph.paragraph_format.keep_with_next = False


def build(template: Path, content_path: Path, output: Path) -> None:
    content = load_content(content_path)
    structure = normalize_structure(content["structure"])
    research_question = normalize_numbered_points(
        content["research_question"], field_name="research_question"
    )
    methods = normalize_numbered_points(content["methods"], field_name="methods")
    results = normalize_numbered_points(content["results"], field_name="results")
    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template, output)

    doc = Document(output)
    validate_template(doc)
    table = doc.tables[0]
    table.autofit = False

    authors = normalize_lines(content["authors"], separator=", ")
    keywords = normalize_lines(content.get("keywords", []), separator="；")

    fill_cell(table.cell(1, 2), content["title"])
    fill_cell(table.cell(2, 2), authors)
    fill_cell(table.cell(2, 4), content["document_type"])
    fill_cell(table.cell(3, 2), content.get("impact_factor") or "—")
    fill_cell(table.cell(3, 4), content["source"])
    fill_cell(table.cell(4, 2), content.get("original_link", ""))
    fill_cell(table.cell(5, 2), keywords)
    fill_cell(table.cell(6, 2), structure, structure_hierarchy=True)
    fill_cell(table.cell(7, 2), research_question)
    fill_cell(table.cell(8, 2), methods)
    fill_cell(table.cell(9, 2), results)
    fill_label(table.cell(10, 1), str(content.get("innovation_label") or "创新与不足"))
    fill_cell(table.cell(10, 2), content["innovation"], bold_headings=True)
    fill_cell(table.cell(11, 2), content["reflection"])
    # The retained remarks row is manual-only. Ignore legacy JSON "notes" values.
    fill_cell(table.cell(12, 1), "")
    compact_trailing_body_paragraphs(doc)

    doc.core_properties.title = str(content["title"])
    doc.core_properties.subject = "文献阅读笔记"
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser(description="Fill the retained literature-reading-note DOCX template.")
    parser.add_argument("--content", required=True, type=Path, help="UTF-8 JSON content file")
    parser.add_argument("--out", required=True, type=Path, help="output DOCX path")
    args = parser.parse_args()

    skill_dir = Path(__file__).resolve().parents[1]
    template = resolve_template(skill_dir)
    if template.resolve() == args.out.resolve():
        raise ValueError("output path must differ from template path")
    build(template.resolve(), args.content.resolve(), args.out.resolve())


if __name__ == "__main__":
    main()
